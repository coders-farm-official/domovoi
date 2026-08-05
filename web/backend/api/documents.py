"""Documents API — the homegrown offline Documents / Spreadsheets / Drawings
suite.

Web-owned feature (no domovoi handler, no ``requires_network`` contract).
Storage is ONE flat ``documents_dir``; the pages are filtered VIEWS by
extension, never separate stores. Three in-app editors, all homegrown or
in-page — the OnlyOffice/Collabora sidecar containers are retired (no
iframes, no JWT/WOPI handshakes, no per-file engine locks):

  * **Doc editor** (web/static/doc_editor.jsx) — markdown-first: toolbar
    formatting + live preview (vendored ``marked``), plain UTF-8 storage
    via the ``/text`` endpoints. Export to .docx server-side
    (``/export/doc``, python-docx).
  * **Sheet editor** (web/static/sheet_editor.jsx) — vendored
    ``x-spreadsheet`` grid with client-side formula evaluation. Reads and
    writes .xlsx (openpyxl — values, formula strings) and .csv through
    the ``/sheet`` endpoints; export to .csv/.xlsx via ``/export/sheet``.
  * **Excalidraw** — in-page React lib, unchanged (.excalidraw scenes).

Legacy office formats (.docx/.doc/.odt/.rtf and .xls/.ods) are stored,
listed, uploaded, and downloaded — but not edited in-app (category
``download``).

EVERY served/saved path is validated inside ``documents_dir`` via the same
realpath / ``relative_to()`` containment check ``music.py`` uses.
"""

from __future__ import annotations

import csv
import io
import logging
import mimetypes
import os
import re
import zipfile
from pathlib import Path
from typing import Any, Literal, Optional

from fastapi import APIRouter, File, HTTPException, Query, Response, UploadFile
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, Field

from domovoi.config import settings as core_settings

log = logging.getLogger(__name__)

router = APIRouter(prefix="/api/documents", tags=["documents"])


# ─── File-type views ────────────────────────────────────────────────
# The Documents ("doc") view is the CATCH-ALL: every file not claimed by
# the Spreadsheets or Drawings views, each routed to the right open-action
# by ``_doc_category`` below. Sheet/drawing views stay narrow (exact-ext).
_DOC_EXTS = frozenset({".docx", ".doc", ".odt", ".rtf", ".txt", ".md"})
_SHEET_EXTS = frozenset({".xlsx", ".xls", ".ods", ".csv"})
_DRAWING_EXTS = frozenset({".excalidraw", ".svg"})

# Legacy office word-processor types — stored/served, not edited in-app.
_OFFICE_WP_EXTS = frozenset({".docx", ".doc", ".odt", ".rtf"})

# Sheet types the homegrown editor can actually round-trip.
_EDITABLE_SHEET_EXTS = frozenset({".xlsx", ".csv"})

# Markdown types the doc editor opens with preview + export.
_MARKDOWN_EXTS = frozenset({".md", ".markdown"})

# Types that open raw in a NEW BROWSER TAB (PDFs + images).
_IMAGE_EXTS = frozenset(
    {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".ico", ".avif", ".svg"}
)
_NEWTAB_EXTS = _IMAGE_EXTS | {".pdf"}

# Known-text extensions the in-app text editor previews confidently. Any
# UNRECOGNIZED type also routes to "text" and is probed via GET /text —
# which falls back gracefully (415 binary/too_large) when it can't decode.
_TEXT_EXTS = frozenset(
    {
        ".txt", ".md", ".markdown", ".rst",
        ".json", ".csv", ".tsv", ".log",
        ".yaml", ".yml", ".ini", ".toml", ".conf", ".cfg", ".env",
        ".xml", ".html", ".htm", ".css",
        ".js", ".jsx", ".ts", ".tsx", ".py", ".sh", ".bash", ".ps1",
        ".sql", ".c", ".h", ".cpp", ".java", ".go", ".rs",
    }
)

_KIND_EXTS: dict[str, frozenset[str]] = {
    "doc": _DOC_EXTS,
    "sheet": _SHEET_EXTS,
    "drawing": _DRAWING_EXTS,
}

# Exts claimed by the other two views — excluded from the doc catch-all.
_NON_DOC_EXTS = _SHEET_EXTS | _DRAWING_EXTS

# Cap on the in-app text editor's read size — bigger than this and we
# tell the UI to fall back to Download/Open-raw rather than stream MBs of
# text into a textarea.
_TEXT_MAX_BYTES = 2 * 1024 * 1024

# Sheet-grid bounds — the homegrown editor is for household spreadsheets,
# not data warehouses. Reads truncate to this window; writes reject beyond.
_SHEET_MAX_ROWS = 1000
_SHEET_MAX_COLS = 60


def _doc_category(ext: str) -> str:
    """How the unified Documents page should open a file of this extension.

      * ``"doc"``      → homegrown markdown editor (.md, toolbar + preview).
      * ``"sheet"``    → homegrown spreadsheet editor (.xlsx/.csv).
      * ``"drawing"``  → in-page Excalidraw editor (.excalidraw scenes).
      * ``"newtab"``   → open the raw endpoint in a new tab (PDFs, images).
      * ``"download"`` → legacy office formats — save-to-device only.
      * ``"text"``     → in-app text editor (known text OR unrecognized;
                         the /text endpoint decides if it's really editable).
    """
    ext = ext.lower()
    if ext in _MARKDOWN_EXTS:
        return "doc"
    if ext in _EDITABLE_SHEET_EXTS:
        return "sheet"
    if ext in _OFFICE_WP_EXTS or ext in (".xls", ".ods"):
        return "download"
    if ext == ".excalidraw":
        return "drawing"
    if ext in _NEWTAB_EXTS:
        return "newtab"
    return "text"


# ─── Containment (reuse of the music.py:349-367 pattern) ────────────
def _documents_dir() -> Path:
    return Path(core_settings.documents_dir).expanduser().resolve(strict=False)


def _safe_target(rel_path: str) -> Path:
    """Resolve ``rel_path`` inside ``documents_dir`` or 400.

    Same realpath / ``relative_to()`` containment check music.py uses.
    Rejects absolute paths and ``..`` traversal: joining an absolute or
    drive-relative path escapes the base, and the post-resolve
    ``relative_to`` catches it.
    """
    if not rel_path or rel_path.strip() == "":
        raise HTTPException(status_code=400, detail="empty document path")
    cleaned = rel_path.replace("\\", "/").lstrip("/")
    base = _documents_dir()
    target = (base / cleaned).resolve(strict=False)
    try:
        target.relative_to(base)
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail=(
                f"refusing path {rel_path!r}: not inside DOCUMENTS_DIR "
                f"({core_settings.documents_dir!r})."
            ),
        )
    return target


def _rel_of(target: Path) -> str:
    """The POSIX-style rel_path (documents_dir-relative) for a resolved
    target."""
    return target.relative_to(_documents_dir()).as_posix()


def _unique_path(dirpath: Path, name: str) -> Path:
    """A non-colliding path inside ``dirpath`` for ``name`` — appends
    `` (1)``, `` (2)`` … to the stem if the target already exists, so a
    second upload of "notes.txt" doesn't clobber the first."""
    target = dirpath / name
    if not target.exists():
        return target
    stem, suffix = target.stem, target.suffix
    i = 1
    while True:
        cand = dirpath / f"{stem} ({i}){suffix}"
        if not cand.exists():
            return cand
        i += 1


# ─── Schemas ────────────────────────────────────────────────────────
class DocumentRow(BaseModel):
    rel_path: str
    name: str
    ext: str
    size: int
    modified_at: float          # epoch seconds
    # Kept for API compatibility (the engine-lock concept is retired —
    # the homegrown editors are last-write-wins like the text editor).
    locked_by: Optional[str] = None
    # How the Documents page opens this file — see ``_doc_category``.
    category: str = "text"


class CreateRequest(BaseModel):
    name: str
    # "text" is the free-form kind: the name is taken VERBATIM — no
    # extension is forced. The others get their default ext appended.
    kind: Literal["doc", "sheet", "drawing", "text"]


class DrawingReadRequest(BaseModel):
    rel_path: str


class DrawingWriteRequest(BaseModel):
    rel_path: str
    content: str
    fmt: Literal["excalidraw", "svg"] = "excalidraw"


class TextWriteRequest(BaseModel):
    text: str


class SheetCell(BaseModel):
    v: Optional[str] = None     # display/stored value
    f: Optional[str] = None     # formula string ("=SUM(A1:A3)") when present


class SheetWriteRequest(BaseModel):
    rows: list[list[Optional[SheetCell]]] = Field(..., max_length=_SHEET_MAX_ROWS)


class UploadResult(BaseModel):
    saved: list[str]
    skipped: list[str]


class DeleteRequest(BaseModel):
    rel_paths: list[str] = Field(..., min_length=1, max_length=1000)


class DeleteResult(BaseModel):
    deleted: list[str]
    failed: list[str]


class ZipRequest(BaseModel):
    rel_paths: list[str] = Field(..., min_length=1, max_length=500)


# ─── List ───────────────────────────────────────────────────────────
@router.get("", response_model=list[DocumentRow])
@router.get("/", response_model=list[DocumentRow])
async def list_documents(
    kind: Literal["all", "doc", "sheet", "drawing"] = Query("all"),
) -> list[DocumentRow]:
    """List files in the flat ``documents_dir`` filtered by extension for
    the requested view, each tagged with the ``category`` that tells the
    UI how to open it."""
    base = _documents_dir()
    exts = _KIND_EXTS.get(kind)
    rows: list[DocumentRow] = []
    if not base.exists():
        return rows

    for entry in sorted(base.iterdir(), key=lambda p: p.name.lower()):
        if not entry.is_file():
            continue
        ext = entry.suffix.lower()
        if kind == "all":
            pass
        elif kind == "doc":
            if ext in _NON_DOC_EXTS:
                continue
        elif ext not in exts:
            continue
        try:
            st = entry.stat()
        except OSError:
            continue
        rows.append(
            DocumentRow(
                rel_path=entry.name,
                name=entry.stem,
                ext=ext,
                size=st.st_size,
                modified_at=st.st_mtime,
                category=_doc_category(ext),
            )
        )
    return rows


# ─── Create a blank file ────────────────────────────────────────────
def _blank_xlsx() -> bytes:
    """A valid empty .xlsx with one blank sheet (openpyxl)."""
    from openpyxl import Workbook

    wb = Workbook()
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _blank_content(ext: str) -> bytes:
    if ext == ".xlsx":
        return _blank_xlsx()
    if ext == ".excalidraw":
        return (
            '{"type":"excalidraw","version":2,"source":"domovoi",'
            '"elements":[],"appState":{},"files":{}}'
        ).encode("utf-8")
    if ext == ".svg":
        return (
            '<svg xmlns="http://www.w3.org/2000/svg" width="800" height="600"></svg>'
        ).encode("utf-8")
    # .md / .txt / .csv — empty file is valid.
    return b""


# Default extension a "New <kind>" gets. "doc" is markdown-first now —
# the homegrown editor's native format.
_KIND_NEW_EXT = {
    "doc": ".md",
    "sheet": ".xlsx",
    "drawing": ".excalidraw",
}


@router.post("/create", response_model=DocumentRow)
async def create_document(req: CreateRequest) -> DocumentRow:
    """Create a new blank file for a view. Name is a bare filename (any
    directory component is stripped). For doc/sheet/drawing the kind's
    default extension is appended when the user didn't type a recognized
    one; the ``text`` kind takes the name VERBATIM."""
    raw = os.path.basename(req.name.replace("\\", "/")).strip()
    if not raw:
        raise HTTPException(status_code=400, detail="empty file name")
    ext = Path(raw).suffix.lower()
    if req.kind != "text" and ext not in _KIND_EXTS[req.kind]:
        ext = _KIND_NEW_EXT[req.kind]
        raw = f"{raw}{ext}"
    target = _safe_target(raw)
    if target.exists():
        raise HTTPException(status_code=409, detail=f"{raw} already exists")
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_bytes(_blank_content(ext))
    st = target.stat()
    return DocumentRow(
        rel_path=target.name,
        name=target.stem,
        ext=ext,
        size=st.st_size,
        modified_at=st.st_mtime,
        category=_doc_category(ext),
    )


# ─── Delete (one or many) ───────────────────────────────────────────
@router.post("/delete", response_model=DeleteResult)
async def delete_documents(req: DeleteRequest) -> DeleteResult:
    """Delete one or more files from ``documents_dir``. Best-effort per
    file — a bad/missing path lands in ``failed`` rather than aborting
    the whole batch."""
    deleted: list[str] = []
    failed: list[str] = []
    for rp in req.rel_paths:
        try:
            target = _safe_target(rp)
        except HTTPException:
            failed.append(rp)
            continue
        try:
            if target.exists() and target.is_file():
                target.unlink()
                deleted.append(_rel_of(target))
            else:
                failed.append(rp)
        except OSError:
            failed.append(rp)
    return DeleteResult(deleted=deleted, failed=failed)


# ─── Bulk download (zip) ────────────────────────────────────────────
@router.post("/download-zip")
async def download_zip(req: ZipRequest) -> Response:
    """Zip the requested files and return the archive (multi-select
    download). Built in memory with an explicit Content-Length; every
    path is containment-checked; missing ones are skipped."""
    buf = io.BytesIO()
    added = 0
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for rp in req.rel_paths:
            try:
                target = _safe_target(rp)
            except HTTPException:
                continue
            if target.exists() and target.is_file():
                z.write(target, arcname=target.name)
                added += 1
    if not added:
        raise HTTPException(status_code=404, detail="none of the requested files exist")
    data = buf.getvalue()
    return Response(
        content=data,
        media_type="application/zip",
        headers={
            "Content-Disposition": 'attachment; filename="documents.zip"',
            "Content-Length": str(len(data)),
        },
    )


# ─── In-app text editor (.txt / .md / unrecognized) ─────────────────
@router.get("/text/{rel_path:path}")
async def read_text_file(rel_path: str) -> Any:
    """Read a file as UTF-8 text for the in-app editors (text + markdown).
    Falls back gracefully instead of choking on non-text input: a file
    over ``_TEXT_MAX_BYTES`` or that isn't valid UTF-8 returns 415 with
    ``{editable:false, reason:"too_large"|"binary"}``."""
    target = _safe_target(rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail=f"{rel_path} not found")
    st = target.stat()
    if st.st_size > _TEXT_MAX_BYTES:
        return JSONResponse(
            status_code=415,
            content={
                "editable": False,
                "reason": "too_large",
                "size": st.st_size,
                "max": _TEXT_MAX_BYTES,
            },
        )
    data = target.read_bytes()
    try:
        body = data.decode("utf-8")
    except UnicodeDecodeError:
        return JSONResponse(
            status_code=415, content={"editable": False, "reason": "binary"}
        )
    newline = "\r\n" if "\r\n" in body else "\n"
    return {
        "rel_path": _rel_of(target),
        "text": body,
        "newline": newline,
        "size": st.st_size,
    }


@router.put("/text/{rel_path:path}", response_model=DocumentRow)
async def write_text_file(rel_path: str, req: TextWriteRequest) -> DocumentRow:
    """Write text back to a file as UTF-8. Containment-checked; newlines
    preserved verbatim."""
    target = _safe_target(rel_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(req.text, encoding="utf-8", newline="")
    st = target.stat()
    ext = target.suffix.lower()
    return DocumentRow(
        rel_path=target.name,
        name=target.stem,
        ext=ext,
        size=st.st_size,
        modified_at=st.st_mtime,
        category=_doc_category(ext),
    )


# ─── Homegrown spreadsheet editor (.xlsx / .csv) ────────────────────
def _read_sheet_grid(target: Path) -> list[list[dict[str, Any]]]:
    """Parse a sheet file into the grid model the editor consumes:
    ``rows[r][c] = {v: str|None, f: str|None}``. Bounded to the
    ``_SHEET_MAX_*`` window (truncated, never errored). .xlsx reads the
    FIRST worksheet; formulas come back as ``f`` with no cached value
    (the client grid re-evaluates)."""
    ext = target.suffix.lower()
    grid: list[list[dict[str, Any]]] = []
    if ext == ".csv":
        with target.open("r", encoding="utf-8-sig", errors="replace", newline="") as fh:
            for r, row in enumerate(csv.reader(fh)):
                if r >= _SHEET_MAX_ROWS:
                    break
                grid.append(
                    [{"v": cell, "f": None} for cell in row[:_SHEET_MAX_COLS]]
                )
        return grid
    if ext == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(target, data_only=False, read_only=True)
        try:
            ws = wb.worksheets[0]
            for r, row in enumerate(ws.iter_rows(max_row=_SHEET_MAX_ROWS,
                                                 max_col=_SHEET_MAX_COLS)):
                cells = []
                for cell in row:
                    val = cell.value
                    if isinstance(val, str) and val.startswith("="):
                        cells.append({"v": None, "f": val})
                    elif val is None:
                        cells.append({"v": None, "f": None})
                    else:
                        cells.append({"v": str(val), "f": None})
                grid.append(cells)
        finally:
            wb.close()
        # Trim fully-empty trailing rows read_only mode can over-report.
        while grid and all(c["v"] is None and c["f"] is None for c in grid[-1]):
            grid.pop()
        return grid
    raise HTTPException(
        status_code=415,
        detail=f"{ext} isn't editable in the sheet editor (xlsx/csv only)",
    )


def _write_sheet_grid(target: Path, rows: list[list[Optional[SheetCell]]]) -> None:
    ext = target.suffix.lower()
    if ext == ".csv":
        with target.open("w", encoding="utf-8", newline="") as fh:
            w = csv.writer(fh)
            for row in rows:
                w.writerow([
                    (c.f if c and c.f else (c.v if c and c.v is not None else ""))
                    for c in row[:_SHEET_MAX_COLS]
                ])
        return
    if ext == ".xlsx":
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        for r, row in enumerate(rows, start=1):
            for c, cell in enumerate(row[:_SHEET_MAX_COLS], start=1):
                if cell is None:
                    continue
                if cell.f:
                    ws.cell(row=r, column=c, value=cell.f)
                elif cell.v is not None and cell.v != "":
                    # Preserve numbers as numbers so exports/formulas work.
                    try:
                        num = float(cell.v)
                        ws.cell(row=r, column=c,
                                value=int(num) if num.is_integer() else num)
                    except ValueError:
                        ws.cell(row=r, column=c, value=cell.v)
        wb.save(target)
        return
    raise HTTPException(
        status_code=415,
        detail=f"{ext} isn't editable in the sheet editor (xlsx/csv only)",
    )


@router.get("/sheet/{rel_path:path}")
async def read_sheet(rel_path: str) -> dict[str, Any]:
    """The grid model for the homegrown sheet editor. 415 for sheet types
    it can't round-trip (.xls/.ods → the UI offers download instead)."""
    target = _safe_target(rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail=f"{rel_path} not found")
    return {
        "rel_path": _rel_of(target),
        "rows": _read_sheet_grid(target),
        "max_rows": _SHEET_MAX_ROWS,
        "max_cols": _SHEET_MAX_COLS,
    }


@router.put("/sheet/{rel_path:path}", response_model=DocumentRow)
async def write_sheet(rel_path: str, req: SheetWriteRequest) -> DocumentRow:
    """Write the editor grid back: .csv gets values (formula strings kept
    verbatim as text), .xlsx gets formulas as formulas and numbers as
    numbers (openpyxl)."""
    target = _safe_target(rel_path)
    if any(len(row) > _SHEET_MAX_COLS for row in req.rows):
        raise HTTPException(status_code=413, detail=f"more than {_SHEET_MAX_COLS} columns")
    target.parent.mkdir(parents=True, exist_ok=True)
    _write_sheet_grid(target, req.rows)
    st = target.stat()
    ext = target.suffix.lower()
    return DocumentRow(
        rel_path=target.name, name=target.stem, ext=ext,
        size=st.st_size, modified_at=st.st_mtime,
        category=_doc_category(ext),
    )


# ─── Exports (docx from markdown · csv/xlsx from sheets) ────────────
_MD_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def _md_runs(paragraph, text_: str) -> None:
    """Split ``text_`` on **bold** / *italic* / `code` markers and add the
    corresponding python-docx runs. Deliberately minimal — the doc editor
    is markdown-first with a basic toolbar, and the export mirrors it."""
    for part in _MD_INLINE.split(text_):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def _markdown_to_docx(text_: str) -> bytes:
    """Markdown → .docx with the structures the editor's toolbar can
    produce: headings, bullet/numbered lists, fenced code blocks, plain
    paragraphs with bold/italic/code inline."""
    from docx import Document

    doc = Document()
    in_code = False
    code_lines: list[str] = []
    for line in text_.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            doc.add_heading(m.group(2), level=min(len(m.group(1)), 6))
            continue
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _md_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _md_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            continue
        if not stripped:
            continue
        _md_runs(doc.add_paragraph(), stripped)
    if in_code and code_lines:  # unterminated fence
        p = doc.add_paragraph()
        p.add_run("\n".join(code_lines)).font.name = "Consolas"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/export/doc/{rel_path:path}")
async def export_doc(rel_path: str, fmt: Literal["docx"] = Query("docx")) -> Response:
    """Export a markdown/text document as .docx (attachment). ``.doc`` is
    a legacy binary format nothing open writes reliably — Word opens
    .docx everywhere the user asked for .doc."""
    target = _safe_target(rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail=f"{rel_path} not found")
    if target.suffix.lower() not in (_MARKDOWN_EXTS | {".txt"}):
        raise HTTPException(status_code=415, detail="docx export is for .md/.txt documents")
    text_ = target.read_bytes().decode("utf-8", errors="replace")
    data = _markdown_to_docx(text_)
    return Response(
        content=data,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": f'attachment; filename="{target.stem}.docx"',
            "Content-Length": str(len(data)),
        },
    )


@router.get("/export/sheet/{rel_path:path}")
async def export_sheet(
    rel_path: str, fmt: Literal["csv", "xlsx"] = Query("csv")
) -> Response:
    """Export a sheet as .csv or .xlsx (attachment), converting through
    the same grid model the editor uses."""
    target = _safe_target(rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail=f"{rel_path} not found")
    grid = _read_sheet_grid(target)  # 415s for non-editable types

    if fmt == "csv":
        out = io.StringIO()
        w = csv.writer(out)
        for row in grid:
            w.writerow([(c["f"] or c["v"] or "") for c in row])
        data = out.getvalue().encode("utf-8")
        media = "text/csv"
        name = f"{target.stem}.csv"
    else:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        for r, row in enumerate(grid, start=1):
            for c, cell in enumerate(row, start=1):
                val = cell["f"] if cell["f"] else cell["v"]
                if val is None or val == "":
                    continue
                if not cell["f"]:
                    try:
                        num = float(val)
                        val = int(num) if num.is_integer() else num
                    except ValueError:
                        pass
                ws.cell(row=r, column=c, value=val)
        buf = io.BytesIO()
        wb.save(buf)
        data = buf.getvalue()
        media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        name = f"{target.stem}.xlsx"
    return Response(
        content=data,
        media_type=media,
        headers={
            "Content-Disposition": f'attachment; filename="{name}"',
            "Content-Length": str(len(data)),
        },
    )


# ─── Upload files into documents_dir ────────────────────────────────
@router.post("/upload", response_model=UploadResult)
async def upload_documents(files: list[UploadFile] = File(...)) -> UploadResult:
    """Upload one or more files straight into ``documents_dir`` from the
    browser. Filenames are sanitized to a bare basename (defanging
    ``../`` / zip-slip traversal across both separators), then deduped."""
    base = _documents_dir()
    try:
        base.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        raise HTTPException(
            status_code=500, detail=f"could not create documents_dir {base}: {e}"
        ) from e

    saved: list[str] = []
    skipped: list[str] = []

    for up in files:
        raw = up.filename or "upload"
        name = os.path.basename(raw.replace("\\", "/")).strip()
        if not name or name in (".", ".."):
            skipped.append(f"{raw!r}: bad filename")
            continue
        data = await up.read()
        target = _unique_path(base, name)
        try:
            target.resolve(strict=False).relative_to(base)
        except ValueError:
            skipped.append(f"{name}: refused (escapes documents_dir)")
            continue
        try:
            target.write_bytes(data)
        except OSError as e:
            skipped.append(f"{name}: write failed ({e})")
            continue
        saved.append(target.name)

    if not saved:
        raise HTTPException(
            status_code=400,
            detail="no files saved" + (f"; {skipped}" if skipped else ""),
        )
    return UploadResult(saved=saved, skipped=skipped)


# ─── Browser-facing raw serve (PDFs / images / open-raw fallback) ───
@router.get("/raw/{rel_path:path}")
async def serve_raw(rel_path: str) -> FileResponse:
    """Serve a file's bytes to the BROWSER (plain ``window.open`` works)
    with the right ``Content-Type`` and an inline ``Content-Disposition``.
    Read scope is bounded entirely by the containment check."""
    target = _safe_target(rel_path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail=f"{rel_path} not found")
    mime, _ = mimetypes.guess_type(target.name)
    return FileResponse(
        str(target),
        media_type=mime or "application/octet-stream",
        filename=target.name,
        content_disposition_type="inline",
    )


# ─── Excalidraw (in-page, no lock) ──────────────────────────────────
@router.post("/drawings/read")
async def read_drawing(req: DrawingReadRequest) -> dict[str, str]:
    """Load an Excalidraw scene (.excalidraw JSON) or .svg for editing."""
    target = _safe_target(req.rel_path)
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"{req.rel_path} not found")
    if target.suffix.lower() not in _DRAWING_EXTS:
        raise HTTPException(status_code=400, detail="not a drawing file")
    return {
        "rel_path": _rel_of(target),
        "content": target.read_text(encoding="utf-8"),
    }


@router.post("/drawings/write")
async def write_drawing(req: DrawingWriteRequest) -> DocumentRow:
    """Save an Excalidraw scene / exported SVG into ``documents_dir``."""
    target = _safe_target(req.rel_path)
    if target.suffix.lower() not in _DRAWING_EXTS:
        raise HTTPException(
            status_code=400, detail="drawing must be .excalidraw or .svg"
        )
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(req.content, encoding="utf-8")
    st = target.stat()
    return DocumentRow(
        rel_path=target.name,
        name=target.stem,
        ext=target.suffix.lower(),
        size=st.st_size,
        modified_at=st.st_mtime,
    )
